from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass
from datetime import date
from typing import Iterable

import httpx
from bs4 import BeautifulSoup

LOGGER = logging.getLogger(__name__)


@dataclass
class RawSegment:
    """Segment of source text before chunking."""

    text: str
    page: int | None = None
    section: str | None = None


@dataclass
class LoadedDocument:
    """Loaded source with normalized metadata and text segments."""

    source_doc: str
    title: str
    doc_type: str
    date: str | None
    segments: list[RawSegment]


class DocumentLoader:
    """Loads documents from multiple scientific sources."""

    def load_pdf(self, file_bytes: bytes, source_name: str) -> LoadedDocument:
        import fitz

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        segments: list[RawSegment] = []
        title = source_name

        metadata = doc.metadata or {}
        if metadata.get("title"):
            title = metadata["title"]

        for page_num, page in enumerate(doc, start=1):
            text = self._normalize_text(page.get_text("text"))
            if not text:
                continue
            segments.append(RawSegment(text=text, page=page_num, section=self._guess_section(text)))

        return LoadedDocument(
            source_doc=source_name,
            title=title,
            doc_type="pdf",
            date=self._safe_date(metadata.get("creationDate")),
            segments=segments,
        )

    def load_docx(self, file_bytes: bytes, source_name: str) -> LoadedDocument:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [self._normalize_text(p.text) for p in doc.paragraphs]
        paragraphs = [p for p in paragraphs if p]

        segments = [RawSegment(text=text, page=None, section=self._guess_section(text)) for text in paragraphs]
        return LoadedDocument(
            source_doc=source_name,
            title=source_name,
            doc_type="docx",
            date=date.today().isoformat(),
            segments=segments,
        )

    def load_url(self, url: str) -> LoadedDocument:
        with httpx.Client(timeout=30.0, follow_redirects=True) as client:
            resp = client.get(url)
            resp.raise_for_status()

        soup = BeautifulSoup(resp.text, "html.parser")
        title = (soup.title.string.strip() if soup.title and soup.title.string else url)

        for tag in soup(["script", "style", "nav", "footer", "aside"]):
            tag.decompose()

        texts = [self._normalize_text(node.get_text(" ")) for node in soup.find_all(["h1", "h2", "h3", "p", "li"])]
        texts = [t for t in texts if t]

        segments: list[RawSegment] = []
        current_section: str | None = None
        for t in texts:
            if len(t.split()) <= 12 and self._looks_like_heading(t):
                current_section = t
                continue
            segments.append(RawSegment(text=t, page=None, section=current_section or self._guess_section(t)))

        return LoadedDocument(
            source_doc=url,
            title=title,
            doc_type="url",
            date=date.today().isoformat(),
            segments=segments,
        )

    def load_arxiv(self, arxiv_id: str) -> LoadedDocument:
        import arxiv

        search = arxiv.Search(id_list=[arxiv_id])
        result = next(search.results(), None)
        if result is None:
            raise ValueError(f"No arXiv record found for ID: {arxiv_id}")

        abstract = self._normalize_text(result.summary or "")
        sections: list[RawSegment] = []
        if abstract:
            sections.append(RawSegment(text=abstract, page=None, section="Abstract"))

        # Optional: attempt to parse PDF full text for better retrieval quality.
        try:
            pdf_bytes = result.download_pdf(dirpath="/tmp", filename=f"{arxiv_id}.pdf")
            with open(pdf_bytes, "rb") as handle:
                loaded = self.load_pdf(handle.read(), source_name=f"arxiv:{arxiv_id}")
            if loaded.segments:
                sections = loaded.segments
        except Exception as exc:  # pragma: no cover - network + filesystem dependent
            LOGGER.warning("Failed to download/parse arXiv PDF for %s: %s", arxiv_id, exc)

        published = result.published.date().isoformat() if result.published else None
        return LoadedDocument(
            source_doc=f"arxiv:{arxiv_id}",
            title=result.title.strip(),
            doc_type="arxiv",
            date=published,
            segments=sections,
        )

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text or "").strip()

    def _guess_section(self, text: str) -> str | None:
        prefixes: Iterable[str] = (
            "abstract",
            "introduction",
            "background",
            "method",
            "methods",
            "results",
            "discussion",
            "conclusion",
            "references",
        )
        lowered = text.lower()
        for p in prefixes:
            if lowered.startswith(p):
                return p.title()
        return None

    def _looks_like_heading(self, text: str) -> bool:
        alpha = [c for c in text if c.isalpha()]
        if not alpha:
            return False
        uppercase_ratio = sum(c.isupper() for c in alpha) / len(alpha)
        return uppercase_ratio > 0.6 or text.istitle()

    def _safe_date(self, raw: str | None) -> str | None:
        if not raw:
            return None
        digits = "".join(c for c in raw if c.isdigit())
        if len(digits) >= 8:
            return f"{digits[:4]}-{digits[4:6]}-{digits[6:8]}"
        return None
